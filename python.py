# investment_app.py

import streamlit as st
import pandas as pd
import numpy as np
from google import genai
from google.genai.errors import APIError
from docx import Document
import io
import json

# --- Cấu hình Trang Streamlit ---
st.set_page_config(
    page_title="App Đánh Giá Phương Án Kinh Doanh",
    layout="wide"
)

st.title("Ứng dụng Đánh giá Phương án Kinh doanh 📈")
st.caption("Sử dụng Gemini AI để trích xuất dữ liệu, tính toán dòng tiền và phân tích hiệu quả dự án.")

# --- Thiết lập API Key ---
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except KeyError:
    st.error("⚠️ **Lỗi Cấu hình:** Vui lòng thêm Khóa API Gemini vào Streamlit Secrets với tên biến `GEMINI_API_KEY`.")
    st.stop()


# --- Hàm đọc file Word ---
@st.cache_data(show_spinner=False)
def read_docx_file(uploaded_file):
    """Đọc nội dung văn bản từ file Word."""
    try:
        # docx.Document() cần đối tượng file-like hỗ trợ seek()
        doc = Document(io.BytesIO(uploaded_file.read()))
        full_text = [para.text for para in doc.paragraphs]
        return "\n".join(full_text)
    except Exception as e:
        return f"Lỗi đọc file Word: {e}"

# --- Hàm gọi API Gemini để trích xuất thông tin (Yêu cầu 1) ---
@st.cache_data
def extract_financial_data(doc_text, api_key):
    """Sử dụng Gemini để trích xuất các thông số tài chính từ văn bản và trả về JSON."""
    
    if not api_key:
        raise ValueError("Khóa API không được cung cấp.")
        
    client = genai.Client(api_key=api_key)
    model_name = 'gemini-2.5-flash'
    
    prompt = f"""
    Bạn là một chuyên gia tài chính và phân tích dự án. Nhiệm vụ của bạn là trích xuất các thông số sau từ nội dung văn bản kinh doanh bên dưới. 
    Các thông số này phải là GIÁ TRỊ SỐ, không có đơn vị (ví dụ: 1000000). 
    
    Vốn đầu tư (Initial Investment - C0): Giá trị tuyệt đối của vốn ban đầu cần bỏ ra.
    Dòng đời dự án (Project Life - N): Số năm hoạt động của dự án.
    WACC (Cost of Capital - k): Tỷ lệ chiết khấu (dạng thập phân, ví dụ: 0.10 cho 10%).
    Thuế suất (Tax Rate - t): Tỷ lệ thuế thu nhập doanh nghiệp (dạng thập phân, ví dụ: 0.20 cho 20%).
    
    Doanh thu hàng năm (Annual Revenue - R): Nếu không có thông tin chi tiết từng năm, hãy ước tính một con số đại diện cho doanh thu hàng năm.
    Chi phí hoạt động hàng năm (Annual Operating Cost - C): Nếu không có thông tin chi tiết từng năm, hãy ước tính một con số đại diện cho chi phí hoạt động hàng năm (chưa bao gồm Khấu hao).
    
    Nếu không tìm thấy thông tin cụ thể, hãy trả về 0 cho giá trị số.
    
    Định dạng đầu ra **bắt buộc** là JSON nguyên mẫu (RAW JSON), không có bất kỳ giải thích hay văn bản nào khác.
    
    {{
      "Vốn đầu tư": <Giá trị số>,
      "Dòng đời dự án": <Giá trị số năm>,
      "Doanh thu hàng năm": <Giá trị số>,
      "Chi phí hoạt động hàng năm": <Giá trị số>,
      "WACC": <Giá trị số thập phân>,
      "Thuế suất": <Giá trị số thập phân>
    }}
    
    Nội dung file Word:
    ---
    {doc_text[:15000]}
    """
    
    try:
        response = client.models.generate_content(
            model=model_name,
            contents=prompt,
            config=genai.types.GenerateContentConfig(
                response_mime_type="application/json",
            )
        )
        # Xử lý chuỗi JSON trả về
        json_str = response.text.strip().replace("```json", "").replace("```", "").strip()
        return pd.Series(json.loads(json_str))
    except APIError as e:
        st.error(f"Lỗi gọi Gemini API: {e}")
        return None
    except json.JSONDecodeError as e:
        st.error(f"Lỗi phân tích JSON từ AI: {e}. Vui lòng thử lại hoặc điều chỉnh file.")
        st.code(response.text if 'response' in locals() else "Không có phản hồi từ AI", language='json')
        return None
    except Exception as e:
        st.error(f"Lỗi không xác định: {e}")
        return None


# --- Hàm xây dựng Dòng tiền (Yêu cầu 2) ---
@st.cache_data(show_spinner=False)
def create_cash_flow_statement(data_series):
    """Xây dựng bảng dòng tiền cơ bản dựa trên dữ liệu trích xuất."""
    
    # Lấy các thông số
    try:
        C0 = data_series["Vốn đầu tư"]
        N = int(data_series["Dòng đời dự án"])
        R = data_series["Doanh thu hàng năm"]
        C = data_series["Chi phí hoạt động hàng năm"]
        t = data_series["Thuế suất"]
    except KeyError as e:
        st.error(f"Dữ liệu trích xuất thiếu: {e}. Vui lòng kiểm tra lại file Word.")
        return None, 0, 0
    
    if N <= 0:
        st.warning("Dòng đời dự án không hợp lệ. Vui lòng kiểm tra dữ liệu trích xuất.")
        return None, 0, 0
    
    # Giả định Khấu hao: Đường thẳng (Straight-Line Depreciation)
    D = C0 / N if N > 0 else 0
    
    # Chuẩn bị DataFrame cho N năm (từ Năm 1 đến Năm N)
    years = list(range(1, N + 1))
    df_cf = pd.DataFrame(index=years)
    
    # Tính toán cho mỗi năm (Giả định dòng tiền đều hàng năm)
    df_cf['Doanh thu (R)'] = R
    df_cf['Chi phí HĐ (C)'] = C
    df_cf['Khấu hao (D)'] = D
    
    # EBT = R - C - D
    df_cf['Lợi nhuận trước thuế (EBT)'] = df_cf['Doanh thu (R)'] - df_cf['Chi phí HĐ (C)'] - df_cf['Khấu hao (D)']
    
    # Tax = EBT * t (chỉ tính thuế khi EBT > 0)
    df_cf['Thuế'] = np.where(df_cf['Lợi nhuận trước thuế (EBT)'] > 0, df_cf['Lợi nhuận trước thuế (EBT)'] * t, 0)
    
    # ATCF (After-Tax Cash Flow) = EBT - Tax + D
    df_cf['Dòng tiền thuần (CF)'] = df_cf['Lợi nhuận trước thuế (EBT)'] - df_cf['Thuế'] + df_cf['Khấu hao (D)']
    
    return df_cf, C0, data_series["WACC"]


# --- Hàm tính toán Chỉ số Tài chính (Yêu cầu 3) ---
def calculate_project_metrics(df_cashflow, initial_investment, wacc):
    """Tính toán NPV, IRR, PP, DPP."""
    
    # Thêm Vốn đầu tư ban đầu vào đầu dòng tiền (Năm 0)
    cash_flows = df_cashflow['Dòng tiền thuần (CF)'].values
    full_cash_flows = np.insert(cash_flows, 0, -initial_investment) 
    
    # 1. NPV
    npv_value = np.npv(wacc, full_cash_flows)
    
    # 2. IRR
    try:
        irr_value = np.irr(full_cash_flows)
    except ValueError:
        irr_value = np.nan
        
    # Chuẩn bị cho PP và DPP
    df_temp = pd.DataFrame({
        'Năm': np.arange(len(full_cash_flows)),
        'CF': full_cash_flows
    })
    
    # 3. PP (Payback Period)
    df_temp['Cum. CF'] = df_temp['CF'].cumsum()
    pp_year = df_temp[df_temp['Cum. CF'] >= 0]['Năm'].min()
    
    if pd.isna(pp_year):
        pp = 'Không hoàn vốn'
    else:
        # Nội suy
        year_before = pp_year - 1
        capital_remaining = abs(df_temp.loc[year_before, 'Cum. CF']) if year_before >= 0 else initial_investment
        cf_of_payback_year = df_temp.loc[pp_year, 'CF']
        pp = year_before + (capital_remaining / cf_of_payback_year) if cf_of_payback_year != 0 else pp_year

    # 4. DPP (Discounted Payback Period)
    discount_factors = 1 / ((1 + wacc) ** df_temp['Năm'])
    df_temp['DCF'] = df_temp['CF'] * discount_factors
    df_temp['Cum. DCF'] = df_temp['DCF'].cumsum()
    dpp_year = df_temp[df_temp['Cum. DCF'] >= 0]['Năm'].min()
    
    if pd.isna(dpp_year):
        dpp = 'Không hoàn vốn'
    else:
        # Nội suy
        year_before_d = dpp_year - 1
        capital_remaining_d = abs(df_temp.loc[year_before_d, 'Cum. DCF']) if year_before_d >= 0 else initial_investment
        dcf_of_payback_year = df_temp.loc[dpp_year, 'DCF']
        dpp = year_before_d + (capital_remaining_d / dcf_of_payback_year) if dcf_of_payback_year != 0 else dpp_year
        
    return npv_value, irr_value, pp, dpp

# --- Hàm gọi AI phân tích chỉ số (Yêu cầu 4) ---
def get_ai_evaluation(metrics_data, wacc_rate, api_key):
    """Gửi các chỉ số đánh giá dự án đến Gemini API và nhận phân tích."""
    
    if not api_key:
        return "Lỗi: Khóa API không được cung cấp."

    try:
        client = genai.Client(api_key=api_key)
        model_name = 'gemini-2.5-flash'
        
        # Định dạng PP và DPP
        pp_str = f"{metrics_data[2]:.2f} năm" if isinstance(metrics_data[2], (int, float)) else metrics_data[2]
        dpp_str = f"{metrics_data[3]:.2f} năm" if isinstance(metrics_data[3], (int, float)) else metrics_data[3]
        
        prompt = f"""
        Bạn là một chuyên gia thẩm định dự án tài chính. Dựa trên các chỉ số hiệu quả của dự án dưới đây, hãy đưa ra nhận xét chi tiết và chuyên sâu, tập trung vào tính khả thi, rủi ro thanh khoản và tính hấp dẫn của dự án.
        
        Các chỉ số cần phân tích:
        1. NPV (Giá trị hiện tại ròng): {metrics_data[0]:,.0f} VND
        2. IRR (Tỷ suất sinh lời nội bộ): {metrics_data[1] * 100:.2f}%
        3. WACC (Tỷ suất chiết khấu): {wacc_rate * 100:.2f}%
        4. PP (Thời gian hoàn vốn): {pp_str}
        5. DPP (Hoàn vốn chiết khấu): {dpp_str}
        
        Yêu cầu phân tích:
        - **Khả thi:** Đánh giá dựa trên nguyên tắc NPV > 0 và IRR > WACC.
        - **Thanh khoản:** Nhận xét về tốc độ thu hồi vốn (PP và DPP). So sánh DPP với dòng đời dự án.
        - **Kết luận:** Tóm tắt ngắn gọn nên *chấp nhận* hay *từ chối* dự án này và lý do chính.
        - Trả lời bằng tiếng Việt, dưới dạng văn xuôi chuyên nghiệp (khoảng 3-4 đoạn).
        """
        response = client.models.generate_content(
            model=model_name,
            contents=prompt
        )
        return response.text
    except APIError as e:
        return f"Lỗi gọi Gemini API: {e}"
    except Exception as e:
        return f"Đã xảy ra lỗi không xác định trong quá trình phân tích: {e}"

# ===============================================
# --- LOGIC CHÍNH CỦA STREAMLIT APP ---
# ===============================================

if 'extracted_data' not in st.session_state: st.session_state['extracted_data'] = None
if 'cash_flow_df' not in st.session_state: st.session_state['cash_flow_df'] = None
if 'metrics' not in st.session_state: st.session_state['metrics'] = None

# --- Khu vực Tải File (Yêu cầu 1) ---
uploaded_file = st.file_uploader(
    "1. Tải file Word (.docx) chứa Phương án Kinh doanh",
    type=['docx']
)

if uploaded_file:
    # --- Nút bấm Lọc Dữ liệu ---
    if st.button("Trích xuất Dữ liệu Tài chính (AI)", key="btn_extract", type="primary", use_container_width=True):
        doc_text = read_docx_file(uploaded_file)
        if doc_text and not doc_text.startswith("Lỗi"):
            with st.spinner('Đang gửi văn bản tới Gemini AI để trích xuất các thông số...'):
                extracted_data_series = extract_financial_data(doc_text, API_KEY)
                st.session_state['extracted_data'] = extracted_data_series
                st.session_state['cash_flow_df'] = None
                st.session_state['metrics'] = None
        elif doc_text.startswith("Lỗi"):
             st.error(doc_text)
                
    if st.session_state['extracted_data'] is not None and not st.session_state['extracted_data'].empty:
        st.divider()
        st.subheader("1. Thông số Tài chính đã Trích xuất (AI) 🤖")
        
        # Hiển thị dữ liệu trích xuất
        df_display = st.session_state['extracted_data'].to_frame('Giá trị').T
        st.dataframe(df_display.style.format({
            'Vốn đầu tư': '{:,.0f}',
            'Doanh thu hàng năm': '{:,.0f}',
            'Chi phí hoạt động hàng năm': '{:,.0f}',
            'WACC': '{:.2%}',
            'Thuế suất': '{:.2%}'
        }), use_container_width=True)

        # --- Xây dựng Dòng tiền và Tính toán Chỉ số (Yêu cầu 2 & 3) ---
        df_cf, C0, WACC = create_cash_flow_statement(st.session_state['extracted_data'])
        
        if df_cf is not None and C0 > 0 and WACC > 0:
            npv, irr, pp, dpp = calculate_project_metrics(df_cf, C0, WACC)
            st.session_state['metrics'] = (npv, irr, pp, dpp)

            # --- Hiển thị Bảng Dòng tiền (Yêu cầu 2) ---
            st.subheader("2. Bảng Dòng tiền Dự án (Năm 1 đến Năm N)")
            st.dataframe(df_cf.style.format({
                'Doanh thu (R)': '{:,.0f}',
                'Chi phí HĐ (C)': '{:,.0f}',
                'Khấu hao (D)': '{:,.0f}',
                'Lợi nhuận trước thuế (EBT)': '{:,.0f}',
                'Thuế': '{:,.0f}',
                'Dòng tiền thuần (CF)': '{:,.0f}'
            }), use_container_width=True)

            # --- Hiển thị Chỉ số Đánh giá (Yêu cầu 3) ---
            st.subheader("3. Các Chỉ số Đánh giá Hiệu quả Dự án")
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.metric(
                    label="NPV (Giá trị hiện tại ròng)", 
                    value=f"{npv:,.0f} VND",
                    delta="Dự án Khả thi" if npv > 0 else "Dự án Không khả thi"
                )
            with col2:
                st.metric(
                    label="IRR (Tỷ suất sinh lời nội bộ)", 
                    value=f"{irr * 100:.2f}%" if not np.isnan(irr) else "N/A",
                    delta="IRR > WACC" if not np.isnan(irr) and irr > WACC else "IRR < WACC hoặc N/A"
                )
            with col3:
                st.metric(label="PP (Thời gian hoàn vốn)", value=f"{pp:.2f} năm" if isinstance(pp, (int, float)) else pp)
            with col4:
                st.metric(label="DPP (Hoàn vốn chiết khấu)", value=f"{dpp:.2f} năm" if isinstance(dpp, (int, float)) else dpp)
                
            st.info(f"Vốn đầu tư ban đầu: **{C0:,.0f} VND** | Tỷ suất chiết khấu (WACC): **{WACC * 100:.2f}%**")
            
            st.divider()

            # --- Phân tích AI (Yêu cầu 4) ---
            st.subheader("4. Phân tích Hiệu quả Dự án (AI) 🧠")
            if st.button("Yêu cầu AI Phân tích Chỉ số", key="btn_analyze_metrics", use_container_width=True):
                with st.spinner('Đang gửi dữ liệu và chờ Gemini phân tích...'):
                    ai_result = get_ai_evaluation(st.session_state['metrics'], WACC, API_KEY)
                    st.session_state['ai_analysis'] = ai_result
            
            if 'ai_analysis' in st.session_state and st.session_state['ai_analysis']:
                st.markdown("**Kết quả Phân tích từ Gemini AI:**")
                st.success(st.session_state['ai_analysis'])

        elif C0 <= 0 or WACC <= 0:
             st.warning("Không thể tính toán: Vốn đầu tư, Dòng đời dự án, hoặc WACC không hợp lệ. Vui lòng kiểm tra dữ liệu trích xuất từ AI.")
    elif st.session_state['extracted_data'] is None:
         st.info("Vui lòng tải file Word lên và bấm nút 'Trích xuất Dữ liệu Tài chính (AI)' để bắt đầu phân tích.")

else:
    st.info("Vui lòng tải lên file Word để bắt đầu đánh giá phương án kinh doanh.")
